"""Phase 1 — make the dry-run harness validate file/format tools.

Covers (per the 2026-06-26 forge/dry-run/stuck-activity fix plan):
  * Task 1.1 — _schema_default_params unwraps a wrapped JSON-Schema via the
    shared schema_utils helper, so the real param names (not the wrapper keys
    "type"/"properties"/"required") drive the defaults.
  * Task 1.2 — _sandbox_paths recognizes source_path and any *_path / *_file /
    *_dir key plus a value heuristic.
  * Task 1.3 — _sandbox_paths writes FORMAT-VALID fixtures by extension
    (docx/xlsx/pdf/png/zip/json, text fallback otherwise).
  * Task 1.4 — an unsynthesizable file format → status="skipped" +
    operator_verify=True (never a hard "failed").
"""
from __future__ import annotations

from pathlib import Path as _P
from unittest.mock import MagicMock, patch

import pytest

from systemu.core.models import Tool, ToolStatus, ToolType


# ─────────────────────────────────────────────────────────────────────────────
# Task 1.1 — _schema_default_params unwraps the wrapped schema

WRAPPED = {
    "type": "object",
    "properties": {
        "source_path": {"type": "string"},
        "password": {"type": "string"},
        "level": {"type": "integer"},
    },
    "required": ["source_path", "password"],
}


def test_wrapped_schema_yields_real_param_names():
    from systemu.pipelines.tool_dry_run import _schema_default_params

    out = _schema_default_params(WRAPPED)
    assert set(out) == {"source_path", "password", "level"}
    assert "properties" not in out and "type" not in out and "required" not in out
    # path + required-string params get a non-empty, sandbox-able placeholder so a
    # forged file tool survives its own required-field check during dry-run.
    assert out["source_path"] and isinstance(out["source_path"], str)
    assert out["password"]            # required string -> non-empty
    assert out["level"] == 0


def test_fallback_params_are_sandboxable_for_a_docx_tool():
    # The exact gap the live tryout caught: the LLM-unavailable fallback used to
    # emit input_path="" -> the tool fails its own "input_path is required" check.
    # Now it emits a non-empty .docx placeholder (inferred from the tool name)
    # that _sandbox_paths materializes into a real file.
    import os
    from systemu.pipelines.tool_dry_run import _schema_default_params, _sandbox_paths
    schema = {"type": "object",
              "properties": {"input_path": {"type": "string"},
                             "output_path": {"type": "string"},
                             "password": {"type": "string"}},
              "required": ["input_path", "output_path", "password"]}
    out = _schema_default_params(schema, tool_name="password_protect_docx")
    assert out["input_path"].endswith(".docx")   # inferred from the tool name
    assert out["password"]                        # required string is non-empty
    sb = _sandbox_paths(dict(out))
    assert os.path.isfile(sb["input_path"])       # a real fixture exists on disk


def test_flat_schema_still_supported():
    from systemu.pipelines.tool_dry_run import _schema_default_params

    assert _schema_default_params({"name": {"type": "string", "default": "Hello"}}) == {"name": "Hello"}


# ─────────────────────────────────────────────────────────────────────────────
# Task 1.2 — _sandbox_paths recognizes source_path + any *_path/*_file key

def test_source_path_key_is_sandboxed():
    from systemu.pipelines.tool_dry_run import _sandbox_paths

    out = _sandbox_paths({"source_path": "C:/real/report.docx"})
    assert out["source_path"] != "C:/real/report.docx"
    assert _P(out["source_path"]).is_file() and out["source_path"].endswith(".docx")


def test_arbitrary_path_suffix_key_sandboxed():
    from systemu.pipelines.tool_dry_run import _sandbox_paths

    out = _sandbox_paths({"report_file": "in/important.pdf", "manifest_path": "x/m.json"})
    assert _P(out["report_file"]).is_file() and _P(out["manifest_path"]).is_file()


def test_non_path_value_left_alone():
    from systemu.pipelines.tool_dry_run import _sandbox_paths

    out = _sandbox_paths({"password": "hunter2", "count": "5"})
    assert out["password"] == "hunter2" and out["count"] == "5"


# ─────────────────────────────────────────────────────────────────────────────
# Task 1.3 — _sandbox_paths writes FORMAT-VALID fixtures by extension

def test_docx_fixture_is_parseable():
    import docx

    from systemu.pipelines.tool_dry_run import _sandbox_paths

    docx.Document(_sandbox_paths({"source_path": "report.docx"})["source_path"])  # no PackageNotFoundError


def test_xlsx_fixture_is_parseable():
    import openpyxl

    from systemu.pipelines.tool_dry_run import _sandbox_paths

    openpyxl.load_workbook(_sandbox_paths({"input_path": "sheet.xlsx"})["input_path"])


def test_pdf_fixture_has_header():
    from systemu.pipelines.tool_dry_run import _sandbox_paths

    assert _P(_sandbox_paths({"file_path": "d.pdf"})["file_path"]).read_bytes().startswith(b"%PDF")


def test_png_fixture_has_magic():
    from systemu.pipelines.tool_dry_run import _sandbox_paths

    assert _P(_sandbox_paths({"image_path": "p.png"})["image_path"]).read_bytes()[:8] == b"\x89PNG\r\n\x1a\n"


def test_zip_fixture_is_valid():
    import zipfile

    from systemu.pipelines.tool_dry_run import _sandbox_paths

    assert zipfile.is_zipfile(_sandbox_paths({"archive_path": "a.zip"})["archive_path"])


def test_unknown_ext_falls_back_to_text():
    from systemu.pipelines.tool_dry_run import _sandbox_paths

    assert _P(_sandbox_paths({"data_path": "x.bin"})["data_path"]).read_bytes()


def test_docx_fixture_is_msoffcrypto_encryptable():
    # END-TO-END: the whole point is that a CORRECT docx-encryption tool passes
    # dry-run. msoffcrypto must be able to OPEN+ENCRYPT the synthesized fixture,
    # not merely have python-docx parse it. If this can't hold for an empty docx,
    # the harness must fall the tool to operator_verify (Task 1.4), never a hard
    # fail — this test pins which path a correct docx tool takes.
    import io

    import msoffcrypto

    from systemu.pipelines.tool_dry_run import _sandbox_paths

    out = _sandbox_paths({"source_path": "report.docx"})
    with open(out["source_path"], "rb") as f:
        of = msoffcrypto.OfficeFile(f)
        buf = io.BytesIO()
        of.encrypt("pw123", buf)        # must not raise on the fixture
    assert buf.getbuffer().nbytes > 0


# ─────────────────────────────────────────────────────────────────────────────
# Task 1.4 — unsynthesizable format → status="skipped" + operator_verify=True

@pytest.fixture
def vault(tmp_path):
    from systemu.vault.vault import Vault

    for sub in ("scrolls", "activities", "shadow_army", "skills", "tools", "evolutions"):
        (tmp_path / sub).mkdir()
        (tmp_path / sub / "index.json").write_text("[]")
    return Vault(str(tmp_path))


def _tool():
    return Tool(
        id="tool_fmt", name="encrypt_docx", description="d",
        tool_type=ToolType.PYTHON_FUNCTION, status=ToolStatus.FORGED, enabled=False,
        implementation_path="vault/tools/impl/encrypt_docx.py",
        parameters_schema={
            "type": "object",
            "properties": {"source_path": {"type": "string"}},
            "required": ["source_path"],
        },
    )


def test_dryrunresult_has_operator_verify_flag():
    from systemu.pipelines.tool_dry_run import DryRunResult

    r = DryRunResult(success=False, status="skipped", operator_verify=True)
    assert r.operator_verify is True and r.to_evidence()["operator_verify"] is True


def test_format_parse_failure_becomes_operator_verify_skip(vault):
    from systemu.pipelines.tool_dry_run import dry_run_tool

    with patch("systemu.pipelines.tool_dry_run._execute",
               return_value={"success": False, "error": "PackageNotFoundError: Package not found"}):
        r = dry_run_tool(_tool(), vault=vault, config=MagicMock(vault_dir="vault"))
    assert r.status == "skipped" and r.operator_verify is True


def test_non_format_failure_still_failed(vault):
    from systemu.pipelines.tool_dry_run import dry_run_tool

    with patch("systemu.pipelines.tool_dry_run._execute",
               return_value={"success": False, "error": "NameError: undefined name 'foo'"}):
        r = dry_run_tool(_tool(), vault=vault, config=MagicMock(vault_dir="vault"))
    assert r.status == "failed" and r.operator_verify is False


# ─────────────────────────────────────────────────────────────────────────────
# Bug B (v0.9.48) — a valid-but-empty/partial LLM param-gen result must backfill
# from schema defaults, never call run() with no args (a FALSE missing-positional).

def _multiparam_tool():
    return Tool(
        id="tool_b", name="password_protect_docx", description="d",
        tool_type=ToolType.PYTHON_FUNCTION, status=ToolStatus.FORGED, enabled=False,
        implementation_path="vault/tools/impl/password_protect_docx.py",
        parameters_schema={
            "type": "object",
            "properties": {"input_path": {"type": "string"},
                           "output_path": {"type": "string"},
                           "password": {"type": "string"}},
            "required": ["input_path", "output_path", "password"],
        },
    )


def test_empty_llm_params_backfills_from_schema():
    from systemu.pipelines.tool_dry_run import _generate_test_params
    with patch("systemu.core.llm_router.llm_call_json", return_value={"params": {}}):
        params, meta = _generate_test_params(_multiparam_tool(), config=MagicMock())
    assert set(params) >= {"input_path", "output_path", "password"}
    assert params["input_path"] and params["password"]      # non-empty, sandbox-able
    assert not meta.get("skip")


def test_bare_empty_dict_llm_result_backfills():
    # raw == {} (no "params" key at all) is the other empty shape.
    from systemu.pipelines.tool_dry_run import _generate_test_params
    with patch("systemu.core.llm_router.llm_call_json", return_value={}):
        params, _ = _generate_test_params(_multiparam_tool(), config=MagicMock())
    assert set(params) >= {"input_path", "output_path", "password"}


def test_partial_llm_params_are_backfilled():
    from systemu.pipelines.tool_dry_run import _generate_test_params
    with patch("systemu.core.llm_router.llm_call_json",
               return_value={"params": {"input_path": "real.docx"}}):
        params, _ = _generate_test_params(_multiparam_tool(), config=MagicMock())
    assert params["input_path"] == "real.docx"               # LLM value wins
    assert params["output_path"] and params["password"]      # gaps backfilled


def test_complete_llm_params_pass_through():
    from systemu.pipelines.tool_dry_run import _generate_test_params
    full = {"input_path": "a.docx", "output_path": "b.docx", "password": "pw"}
    with patch("systemu.core.llm_router.llm_call_json", return_value={"params": dict(full)}):
        params, _ = _generate_test_params(_multiparam_tool(), config=MagicMock())
    assert params == full


def test_skip_advice_still_honored():
    from systemu.pipelines.tool_dry_run import _generate_test_params
    with patch("systemu.core.llm_router.llm_call_json",
               return_value={"skip_dry_run": True, "skip_reason": "destructive"}):
        params, meta = _generate_test_params(_multiparam_tool(), config=MagicMock())
    assert meta.get("skip") is True


# ── v0.9.51 graceful degradation: failure on an UN-synthesizable constrained
#    param routes to operator_verify, not a doomed "failed" ─────────────────────

def _pattern_tool():
    return Tool(
        id="tool_code", name="make_code", description="d",
        tool_type=ToolType.PYTHON_FUNCTION, status=ToolStatus.FORGED, enabled=False,
        implementation_path="vault/tools/impl/make_code.py",
        parameters_schema={
            "type": "object",
            "properties": {"code": {"type": "string", "pattern": "^[A-Z]{3}-[0-9]{4}$"}},
            "required": ["code"],
        },
    )


def test_unresolved_constrained_param_failure_becomes_operator_verify(vault):
    from systemu.pipelines.tool_dry_run import dry_run_tool
    with patch("systemu.pipelines.tool_dry_run._execute",
               return_value={"success": False, "error": "ValueError: code must match pattern"}):
        r = dry_run_tool(_pattern_tool(), vault=vault, config=MagicMock(vault_dir="vault"))
    assert r.status == "skipped" and r.operator_verify is True


# ── v0.9.51 context-grounding: forge capture + dry-run read ───────────────────

def test_capture_grounding_sets_tool_inputs_from_scroll():
    from systemu.pipelines.tool_forge import _capture_grounding
    from systemu.core.models import Scroll
    t = _tool()
    scroll = Scroll(id="s1", name="n", source_session_id="x",
                    raw_instructions_path="", narrative_md="",
                    raw_request="please password-protect report.docx")
    _capture_grounding(t, scroll)
    assert t.grounding_inputs == ["report.docx"]


def test_capture_grounding_noop_without_scroll():
    from systemu.pipelines.tool_forge import _capture_grounding
    t = _tool()
    _capture_grounding(t, None)
    assert t.grounding_inputs == []


def test_dry_run_grounds_params_from_tool_inputs(vault, tmp_path):
    from systemu.pipelines.tool_dry_run import dry_run_tool
    import docx
    real = tmp_path / "report.docx"
    d = docx.Document(); d.add_paragraph("REAL CONTENT"); d.save(str(real))
    t = _tool()                          # name=encrypt_docx, param source_path
    t.grounding_inputs = [str(real)]
    captured = {}

    def _fake_execute(tool, params, **kw):
        captured["params"] = params
        return {"success": True, "parsed": {}}

    with patch("systemu.pipelines.tool_dry_run._execute", _fake_execute):
        dry_run_tool(t, vault=vault, config=MagicMock(vault_dir="vault"))
    used = captured["params"]["source_path"]
    assert "REAL CONTENT" in docx.Document(used).paragraphs[0].text   # real, grounded
