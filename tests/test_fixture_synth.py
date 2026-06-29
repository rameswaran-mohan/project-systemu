"""v0.9.51 — exhaustive shape tests for the schema-walk fixture synthesizer.

One test per failure-space shape (scalar path, list-of-paths, nested-object path,
map-of-paths, dir-vs-file, multi-format, formats, enum/const, bounds, wrapped
schema, anyOf/allOf, $ref+cycle, optional-vs-required). The contract: the params
are structurally complete and every path leaf is a REAL fixture file on disk.
"""
from __future__ import annotations

import os

import pytest

from systemu.pipelines.fixture_synth import synthesize_params, looks_like_path


def _synth(schema, tmp_path, tool_name=""):
    return synthesize_params(schema, tool_name=tool_name, sandbox_dir=str(tmp_path))


# ── scalar path → real format-valid fixture ───────────────────────────────────

def test_scalar_path_materializes_real_file(tmp_path):
    r = _synth({"type": "object", "properties": {"input_path": {"type": "string"}},
                "required": ["input_path"]}, tmp_path, tool_name="password_protect_docx")
    p = r.params["input_path"]
    assert os.path.isfile(p) and p.endswith(".docx")        # ext inferred from tool name
    import docx; docx.Document(p)                            # parseable, not a text blob


def test_wrapped_schema_unwraps(tmp_path):
    # the original bug: a wrapped {type,properties,required} must not leak wrapper keys
    r = _synth({"type": "object",
                "properties": {"output_path": {"type": "string"}, "password": {"type": "string"}},
                "required": ["output_path", "password"]}, tmp_path)
    assert set(r.params) == {"output_path", "password"}
    assert "properties" not in r.params and "type" not in r.params


def test_flat_schema_also_supported(tmp_path):
    r = _synth({"input_path": {"type": "string"}, "n": {"type": "integer"}}, tmp_path)
    assert set(r.params) == {"input_path", "n"} and r.params["n"] == 0


# ── list of paths (the v0.9.49 bug) → N distinct real fixtures ────────────────

def test_list_of_paths_materializes_each_item(tmp_path):
    r = _synth({"type": "object",
                "properties": {"files_to_add": {"type": "array", "items": {"type": "string"},
                                                "minItems": 2}},
                "required": ["files_to_add"]}, tmp_path)
    lst = r.params["files_to_add"]
    assert isinstance(lst, list) and len(lst) == 2
    assert all(os.path.isfile(x) for x in lst)              # every item is a real file
    assert lst[0] != lst[1]                                  # distinct fixtures


# ── nested object with a path → materialized at the nested leaf ───────────────

def test_nested_object_path_is_materialized(tmp_path):
    r = _synth({"type": "object", "properties": {
        "config": {"type": "object", "properties": {"output_file": {"type": "string"}},
                   "required": ["output_file"]}}}, tmp_path)
    inner = r.params["config"]["output_file"]
    assert os.path.isfile(inner)


def test_array_of_objects_with_path(tmp_path):
    r = _synth({"type": "object", "properties": {
        "jobs": {"type": "array", "minItems": 1,
                 "items": {"type": "object", "properties": {"src_path": {"type": "string"}},
                           "required": ["src_path"]}}}}, tmp_path)
    assert os.path.isfile(r.params["jobs"][0]["src_path"])


def test_additional_properties_map(tmp_path):
    r = _synth({"type": "object", "properties": {},
                "additionalProperties": {"type": "string", "format": "path"}}, tmp_path)
    # one representative entry synthesized through the value subschema
    assert any(isinstance(v, str) and os.path.isfile(v) for v in r.params.values())


# ── dir vs file ───────────────────────────────────────────────────────────────

def test_dir_key_creates_directory(tmp_path):
    r = _synth({"type": "object", "properties": {"output_dir": {"type": "string"}},
                "required": ["output_dir"]}, tmp_path)
    assert os.path.isdir(r.params["output_dir"])


# ── multiple formats in one call → each gets the right valid bytes ────────────

def test_multi_format_each_correct(tmp_path):
    r = _synth({"type": "object", "properties": {
        "report_docx": {"type": "string"}, "sheet_xlsx": {"type": "string"},
        "logo_png": {"type": "string"}}}, tmp_path)
    assert r.params["report_docx"].endswith(".docx")
    assert r.params["sheet_xlsx"].endswith(".xlsx")
    assert r.params["logo_png"].endswith(".png")
    assert os.path.isfile(r.params["logo_png"])
    assert open(r.params["logo_png"], "rb").read()[:8] == b"\x89PNG\r\n\x1a\n"


# ── format-constrained scalars, enum/const, numeric bounds, required ──────────

def test_string_formats(tmp_path):
    r = _synth({"type": "object", "properties": {
        "to": {"type": "string", "format": "email"},
        "when": {"type": "string", "format": "date"},
        "id": {"type": "string", "format": "uuid"}}}, tmp_path)
    assert "@" in r.params["to"] and r.params["when"] == "2020-01-01"
    assert len(r.params["id"]) == 36


def test_enum_and_const(tmp_path):
    r = _synth({"type": "object", "properties": {
        "mode": {"enum": ["fast", "slow"]}, "v": {"const": 7}}}, tmp_path)
    assert r.params["mode"] == "fast" and r.params["v"] == 7


def test_numeric_bounds(tmp_path):
    r = _synth({"type": "object", "properties": {
        "level": {"type": "integer", "minimum": 3, "maximum": 9}}}, tmp_path)
    assert r.params["level"] == 3


def test_required_string_is_nonempty(tmp_path):
    r = _synth({"type": "object", "properties": {"password": {"type": "string"}},
                "required": ["password"]}, tmp_path)
    assert r.params["password"]                              # non-empty for a required field


# ── composition: anyOf / allOf / $ref + cycle ─────────────────────────────────

def test_anyof_picks_synthesizable_branch(tmp_path):
    r = _synth({"type": "object", "properties": {
        "x": {"anyOf": [{"type": "integer"}, {"type": "string"}]}}}, tmp_path)
    assert r.params["x"] in (0,) or isinstance(r.params["x"], (int, str))


def test_allof_merges(tmp_path):
    r = _synth({"type": "object", "properties": {
        "p": {"allOf": [{"type": "string"}, {"minLength": 8}]}}}, tmp_path)
    assert isinstance(r.params["p"], str) and len(r.params["p"]) >= 8


def test_ref_and_cycle_do_not_hang(tmp_path):
    schema = {"type": "object", "$defs": {"node": {"$ref": "#/$defs/node"}},
              "properties": {"n": {"$ref": "#/$defs/node"}, "ok": {"type": "string"}}}
    r = _synth(schema, tmp_path)                             # must terminate (cycle guard)
    assert "ok" in r.params


# ── the path oracle ───────────────────────────────────────────────────────────

def test_oracle_signals():
    assert looks_like_path("input_path", {"type": "string"})[0] is True
    assert looks_like_path("x", {"type": "string", "format": "path"})[0] is True
    assert looks_like_path("report", {"description": "the .docx file to read"})[0] is True
    assert looks_like_path("out_dir", {"type": "string"})[1] == "dir"
    assert looks_like_path("password", {"type": "string"})[0] is False


def test_created_paths_recorded(tmp_path):
    r = _synth({"type": "object", "properties": {"a_path": {"type": "string"}}}, tmp_path)
    assert r.params["a_path"] in r.created_paths


# ── graceful degradation: a constraint the engine can't satisfy is FLAGGED ────

def test_pattern_constrained_required_string_flagged_unresolved(tmp_path):
    # no regex engine — a `pattern`-bound required string can't be guaranteed, so
    # it's flagged (the caller degrades a downstream failure to operator_verify
    # instead of a doomed hard-fail).
    r = _synth({"type": "object", "properties": {
        "code": {"type": "string", "pattern": "^[A-Z]{3}-[0-9]{4}$"}},
        "required": ["code"]}, tmp_path)
    assert "code" in r.unresolved
    assert r.params["code"]                  # still a best-effort non-empty value


def test_unconstrained_string_not_flagged(tmp_path):
    r = _synth({"type": "object", "properties": {"name": {"type": "string"}},
                "required": ["name"]}, tmp_path)
    assert r.unresolved == []


# ── context-grounding: real operator files → real dry-run inputs ──────────────

def _real_docx(p, text):
    import docx
    d = docx.Document(); d.add_paragraph(text); d.save(str(p))
    return str(p)


def test_grounding_path_used_for_input_leaf(tmp_path):
    real = _real_docx(tmp_path / "report.docx", "REAL CONTENT")
    sb = tmp_path / "sb"
    r = synthesize_params(
        {"type": "object", "properties": {"input_path": {"type": "string"}},
         "required": ["input_path"]},
        tool_name="protect_docx", sandbox_dir=str(sb), grounding_paths=[real])
    used = r.params["input_path"]
    import docx
    assert "REAL CONTENT" in docx.Document(used).paragraphs[0].text   # real content, not empty
    assert os.path.dirname(used) == str(sb)                            # a COPY in the sandbox
    assert used != real                                               # not the original


def test_grounding_missing_file_falls_back_to_synthetic(tmp_path):
    sb = tmp_path / "sb"
    r = synthesize_params(
        {"type": "object", "properties": {"input_path": {"type": "string"}}},
        tool_name="protect_docx", sandbox_dir=str(sb),
        grounding_paths=["C:/nope/ghost.docx"])           # doesn't exist → ignored
    assert os.path.isfile(r.params["input_path"])         # synthetic fixture, no crash


def test_grounding_not_burned_on_output_leaf(tmp_path):
    real = _real_docx(tmp_path / "src.docx", "REAL CONTENT")
    sb = tmp_path / "sb"
    r = synthesize_params(
        {"type": "object", "properties": {
            "input_path": {"type": "string"}, "output_path": {"type": "string"}},
         "required": ["input_path", "output_path"]},
        tool_name="t", sandbox_dir=str(sb), grounding_paths=[real])
    import docx
    assert "REAL CONTENT" in docx.Document(r.params["input_path"]).paragraphs[0].text
    # output is written, not read → synthetic, NOT a copy of the grounding file
    assert b"REAL CONTENT" not in open(r.params["output_path"], "rb").read()
    assert r.params["output_path"] != r.params["input_path"]


def test_extract_candidate_paths():
    from systemu.pipelines.fixture_synth import extract_candidate_paths
    got = extract_candidate_paths("please password-protect report.docx and merge data.xlsx",
                                  "save to C:/Users/me/out.pdf")
    assert "report.docx" in got and "data.xlsx" in got
    assert any(g.endswith("out.pdf") for g in got)
    assert extract_candidate_paths("no files here") == []


def test_flat_schema_no_required_fills_all_strings_nonempty(tmp_path):
    # the create_password_protected_zip bug: a flat schema with NO `required` list →
    # every declared string must still be non-empty (run() does `if not password`).
    r = _synth({"output_path": {"type": "string"},
                "files_to_add": {"type": "array", "items": {"type": "string"}},
                "password": {"type": "string"}}, tmp_path)
    assert r.params["password"] == "dryrun"
    assert r.params["files_to_add"]


def test_explicit_required_list_keeps_optional_empty(tmp_path):
    r = _synth({"type": "object", "properties": {"a": {"type": "string"}, "b": {"type": "string"}},
                "required": ["a"]}, tmp_path)
    assert r.params["a"] == "dryrun" and r.params["b"] == ""
