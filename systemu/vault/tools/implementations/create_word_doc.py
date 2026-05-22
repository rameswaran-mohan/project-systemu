#!/usr/bin/env python3
"""Create a .docx Word document with optional title, body text, and image."""
from __future__ import annotations

import re
from pathlib import Path

TOOL_META = {
    "name": "create_word_doc",
    "tool_type": "file",
    "dependencies": ["python-docx"],
}


# schema mirrors the vault JSON definition so callers can introspect
# the tool contract directly from the module.  The output_path description is
# the operator-facing contract and MUST stay in sync with
# systemu/vault/tools/tool_tool_6e6e62c0.json.
TOOL_SPEC = {
    "name": "create_word_doc",
    "description": (
        "Create a Word .docx document with title, body text, sections, and "
        "an optional embedded image."
    ),
    "tool_type": "file_operation",
    "parameters_schema": {
        "type": "object",
        "properties": {
            "output_path": {
                "type": "string",
                "description": (
                    "Absolute path where the .docx should be written.  May be "
                    "either a full file path (e.g. '/app/outputs/report.docx') "
                    "or a directory (e.g. '/app/outputs/') — if a directory is "
                    "given, a filename will be derived from the document title."
                ),
            },
            "title": {
                "type": "string",
                "description": "Document heading",
                "default": "",
            },
            "body_text": {
                "type": "string",
                "description": "Body paragraph text",
                "default": "",
            },
            "sections": {
                "type": "array",
                "description": (
                    "Optional list of {heading, content} section objects "
                    "rendered after the title."
                ),
                "default": [],
            },
            "image_path": {
                "type": "string",
                "description": "Path to image to embed",
                "default": "",
            },
            "overwrite": {
                "type": "boolean",
                "description": "Overwrite if exists",
                "default": True,
            },
        },
        "required": ["output_path"],
    },
}


def _slugify(text: str) -> str:
    """Turn a doc title into a filesystem-safe slug for use as a filename."""
    slug = re.sub(r"[^\w\s-]", "", text or "").strip()
    slug = re.sub(r"[-\s]+", "_", slug)
    return slug or "document"


def _resolve_output_file(output_path: str, title: str) -> Path:
    """turn an output_path that might be a directory, a path without
    extension, or a path with a trailing separator, into a concrete .docx
    file path under the intended directory."""
    op = Path(output_path).expanduser()
    if op.exists() and op.is_dir():
        return op / f"{_slugify(title)}.docx"
    if output_path.endswith(("/", "\\")):
        op.mkdir(parents=True, exist_ok=True)
        return op / f"{_slugify(title)}.docx"
    if op.suffix == "":
        # No extension and not an existing dir — assume the LLM meant a file
        # but forgot .docx
        return op.with_suffix(".docx")
    return op


def run(**kwargs) -> dict:
    output_path: str = kwargs.get("output_path", "")
    title: str = kwargs.get("title", "")
    body_text: str = kwargs.get("body_text", "")
    sections = kwargs.get("sections") or []
    image_path: str = kwargs.get("image_path", "")
    overwrite: bool = bool(kwargs.get("overwrite", True))

    if not output_path:
        return {"success": False, "output_path": "", "error": "output_path is required"}

    try:
        from docx import Document
        from docx.shared import Inches

        # resolve dir / extensionless / trailing-separator paths
        # gracefully so the LLM doesn't have to know whether it was handed a
        # directory or a file.
        target_path = _resolve_output_file(output_path, title)

        if target_path.exists() and not overwrite:
            return {
                "success": False,
                "output_path": str(target_path),
                "error": "File already exists",
            }

        target_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        if title:
            doc.add_heading(title, 0)

        if body_text:
            doc.add_paragraph(body_text)

        # structured sections support — each entry is {heading, content}
        for section in sections:
            if not isinstance(section, dict):
                continue
            heading = section.get("heading")
            content = section.get("content")
            if heading:
                doc.add_heading(str(heading), level=1)
            if content:
                doc.add_paragraph(str(content))

        if image_path:
            img = Path(image_path).expanduser()
            if img.exists():
                doc.add_picture(str(img), width=Inches(6))

        doc.save(str(target_path))

        return {"success": True, "output_path": str(target_path), "error": None}

    except Exception as exc:
        return {"success": False, "output_path": "", "error": str(exc)}


# callable alias so callers can import the tool by name without going
# through the dispatch dict.  Mirrors the run(**kwargs) signature.
def create_word_doc(**kwargs) -> dict:
    return run(**kwargs)
